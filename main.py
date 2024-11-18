"""
Программа для подсчета частоты слов в документе word
"""

import os
import docx
import re
from tkinter import filedialog as fd

def selectFile(fileName):
    path = fd.askdirectory() + '/' + fileName
    return path.replace('/', '\\')


#mydoc = docx.Document()
#Исключаем предлоги
wordNot = ['на','под','за','к','из','по','об','от','в',
           'у','с','о','над','около','при','перед','также',
           'что', 'и', 'для','после','его']
fileName = 'Образ_совместного_решения_черновик.docx'
TEXT_FILE = selectFile(fileName)
if TEXT_FILE:
    doc = docx.Document(TEXT_FILE)
    result = [p.text for p in doc.paragraphs]
    resultStr=''
    for item in result:
        resultStr += " " + item
    resultStr = resultStr.lower()
    match_pattern = re.findall(r'[А-я]+', resultStr)
    frequency = {}

    for word in match_pattern:
        if word not in wordNot:
            count = frequency.get(word, 0)
            frequency[word] = count + 1


    sorted_dict={}
    sorted_keys = sorted(frequency, key=frequency.get)  # [1, 3, 2]

    for w in sorted_keys:
        sorted_dict[w] = frequency[w]

    frequency_list = sorted_dict.keys()

    for words in frequency_list:
        print(words, sorted_dict[words])
    g = 0